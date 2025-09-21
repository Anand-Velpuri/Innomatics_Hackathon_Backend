import os
import fitz  # PyMuPDF
import docx
import json
from fuzzywuzzy import fuzz
# from sentence_transformers import SentenceTransformer, util
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from pydantic import BaseModel, Field
from langchain_core.output_parsers import JsonOutputParser
from typing import Optional, List
import os

# --- LLM and Model Configuration ---
# It's best practice to load models and clients once
try:
    llm = ChatOpenAI(
        model=os.environ.get("MODEL_NAME"),
        temperature=0,
        api_key=os.environ.get("CUSTOM_API_KEY"),
        base_url=os.environ.get("CUSTOM_BASE_URL"),
    )
    # sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
    print("AI models loaded successfully.")
except Exception as e:
    print(f"Error loading AI models: {e}")
    llm = None
    # sentence_model = None

# --- Re-using the Pydantic class for JD Parsing ---
class JobDescription(BaseModel):
    job_title: str = Field(description="The title of the job role.")
    must_have_skills: list[str] = Field(description="A list of essential skills.")
    good_to_have_skills: list[str] = Field(description="A list of preferred skills.")

# --- New function to parse job document for pre-filling form ---
class JobDocument(BaseModel):
    job_title: str = Field(description="The extracted job title from the document.")
    department: str = Field(description="The department or team for the job role.")
    description: str = Field(description="A detailed summary or description of the job role.")
    requirements: str = Field(description="A list of key skills, qualifications, or experience required for the job.")

# NEW: Create a wrapper model for a LIST of jobs
class JobDocumentList(BaseModel):
    jobs: List[JobDocument] = Field(description="A list of all job postings extracted from the document.")

def parse_job_document_to_fill_form(doc_text: str) -> dict:
    """Uses an LLM to parse a job description document into a list of structured job objects."""
    # Use the new JobDocumentList as the parser's target
    parser = JsonOutputParser(pydantic_object=JobDocumentList)
    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are an expert HR assistant. Extract all job postings from the document into a structured JSON list. Respond ONLY with the JSON object."),
        ("human", "Here is the document text:\n\n{doc}\n\n{format_instructions}")
    ])
    chain = prompt | llm | parser
    # The result will be a dict like {"jobs": [...]}, so we return the list inside.
    return chain.invoke({"doc": doc_text, "format_instructions": parser.get_format_instructions()})

# --- All the analysis functions from the notebook go here ---

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    # (Same function as in Colab)
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = "".join(page.get_text() for page in doc)
    doc.close()
    return text

def extract_text_from_docx(docx_bytes: bytes) -> str:
    # (Same function as in Colab, but needs io)
    import io
    document = docx.Document(io.BytesIO(docx_bytes))
    return "\n".join([para.text for para in document.paragraphs])

def parse_job_description(jd_text: str) -> dict:
    parser = JsonOutputParser(pydantic_object=JobDescription)
    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are an expert HR assistant. Extract structured information from a job description. Respond ONLY with the requested JSON object."),
        ("human", "Job Description:\n\n{jd}\n\n{format_instructions}")
    ])
    chain = prompt | llm | parser
    return chain.invoke({"jd": jd_text, "format_instructions": parser.get_format_instructions()})

def calculate_hard_match(resume_text: str, parsed_jd: dict) -> dict:
    # (Same logic as in Colab)
    must_have_skills = parsed_jd.get("must_have_skills", [])
    good_to_have_skills = parsed_jd.get("good_to_have_skills", [])
    found_skills, missing_skills = [], []
    for skill in must_have_skills:
        if fuzz.token_set_ratio(resume_text.lower(), skill.lower()) > 80:
            found_skills.append(skill)
        else:
            missing_skills.append(skill)
    must_have_score = len(found_skills) / len(must_have_skills) if must_have_skills else 1.0
    return {"score": must_have_score, "missing_skills": missing_skills}

# def calculate_semantic_match(resume_text: str, jd_text: str) -> float:
#     # (Same logic as in Colab)
#     resume_embedding = sentence_model.encode(resume_text, convert_to_tensor=True)
#     jd_embedding = sentence_model.encode(jd_text, convert_to_tensor=True)
#     cosine_score = util.cos_sim(resume_embedding, jd_embedding)
#     return cosine_score.item()

def calculate_semantic_match(resume_text: str, jd_text: str) -> float:
    """
    Calculates the semantic similarity using TF-IDF and cosine similarity
    from scikit-learn.
    """
    vectorizer = TfidfVectorizer(stop_words='english')
    corpus = [resume_text, jd_text]
    
    # Generate TF-IDF vectors
    tfidf_matrix = vectorizer.fit_transform(corpus)
    
    # Calculate cosine similarity between the two vectors
    # The result is a matrix, so we access the specific score at [0, 1]
    similarity_score = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
    
    return float(similarity_score)


def parse_job_document_to_fill_form(doc_text: str) -> dict:
    """Uses an LLM to parse a job description document into structured fields for a form."""
    parser = JsonOutputParser(pydantic_object=JobDocument)
    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are an expert HR assistant. Your task is to extract key details from a job description document. Respond ONLY with the JSON object."),
        ("human", "Here is the document text:\n\n{doc}\n\n{format_instructions}")
    ])
    chain = prompt | llm | parser
    return chain.invoke({"doc": doc_text, "format_instructions": parser.get_format_instructions()})


def get_final_verdict_and_feedback(hard_match_res: dict, semantic_score: float, parsed_jd: dict) -> dict:
    # (Combined logic from Colab)
    hard_match_score = hard_match_res['score']
    final_score = (hard_match_score * 0.65) + (semantic_score * 0.35)
    final_score_percentage = round(final_score * 100)
    
    if final_score_percentage >= 80: verdict = "High Suitability"
    elif 60 <= final_score_percentage < 80: verdict = "Medium Suitability"
    else: verdict = "Low Suitability"
        
    missing_skills_str = ", ".join(hard_match_res['missing_skills']) or "None"
    feedback_prompt = f"""
    As a career coach, provide brief, constructive feedback for a resume submitted for the role of '{parsed_jd.get('job_title', 'N/A')}'.
    The resume scored {final_score_percentage}% ({verdict}).
    The following skills were missing or unclear: {missing_skills_str}.
    Suggest one key improvement. Keep it concise and encouraging.
    """
    feedback_text = llm.invoke(feedback_prompt).content
    
    return {
        "job_title": parsed_jd.get('job_title'),
        "relevance_score": final_score_percentage,
        "verdict": verdict,
        "missing_skills": hard_match_res['missing_skills'],
        "feedback": feedback_text
    }
